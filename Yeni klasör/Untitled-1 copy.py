import pyperclip
import re
import win32com.client
from docx import Document

# def kopyala_ve_devam_et(dosya_adi):
#     with open(dosya_adi, 'r', encoding='UTF-8') as dosya:
#         tam_metin = dosya.read()

#         # Kelimeleri ayır
#         kelimeler = tam_metin.split()

#         # İlk 920 kelimeyi al
#         kisilmis_metin = ' '.join(kelimeler[:920])

#         # Son cümlenin sonuna kadar al
#         son_cumle = ' '.join(re.split(r'(?<=\.)\s', kisilmis_metin)[-1:])
        
#         # Kırpılmış metni ve son cümleyi panoya kopyala
#         pyperclip.copy(kisilmis_metin + son_cumle)

# dosya_adi = "C:/Dev/belge.docx"  # Dosya adını değiştirin
# kopyala_ve_devam_et(dosya_adi)
# def word2string(dosya):
#     word=win32com.client.Dispatch("Word.Application")
#     word.visible=False
#     wb=word.Documents.Open(dosya)
#     doc=word.ActiveDocument
#     text=doc.Range().Text
#     doc.Close(False)
#     # word.Quit()
#     return text

# def metni_kes(text):
#     kelimeler = text.split()
#     if len(kelimeler)<920:
#         print("920 den az kelime var!!!")
#     # İlk 920 kelimeyi al
#     kisilmis_metin = ' '.join(kelimeler[:920])

#     # Son cümlenin sonuna kadar al
#     son_cumle = ' '.join(re.split(r'(?<=\.)\s', kisilmis_metin)[-1:])
    
#     i=920
#     while(kelimeler[i][-1]!='.'):
#         i=i+1
#     print(kelimeler[i])

#     kisilmis_metin = ' '.join(kelimeler[:i+1])
#     # Kırpılmış metni ve son cümleyi panoya kopyala
#     pyperclip.copy(kisilmis_metin)
    
# text=word2string("C:/Dev/belge.docx")
# metni_kes(text)



# Word belgesini aç
doc = Document('C:/Dev/belge.docx')

# Belgedeki paragrafları ve başlıkları döngü ile oku
for paragraph in doc.paragraphs:
    print(paragraph.text)  # Paragrafı yazdır

# Başlıkları almak için ayrı bir döngü oluşturabilirsiniz
for paragraph in doc.paragraphs:
    if paragraph.style.name.startswith('Heading'):
        print("Başlık:", paragraph.text)  # Başlığı yazdır


